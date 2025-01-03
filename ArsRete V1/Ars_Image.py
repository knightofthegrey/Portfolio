#ArsImage
#Image processing module designed to implement Ars_Network for image processing
#The initial task of this program is to generate samples of typewritten text, convert to images, and test training Ars_Network to identify them.

import Ars_Network
import docx
import docx2pdf
import fitz
import os
import random
import numpy as np
from PIL import Image
from PIL import ImageDraw
import math
from time import time

#__GLOBAL VARIABLES__
global_img_dir = "ArsImg"

#__GENERAL IMAGE PROCESSING FUNCTIONS__

def pad(in_img,pad_x,pad_y):
    #Pad image with white pixels to a given size so they can be inputted into the MLP network
    #The Network object requires inputted images to be of uniform size.
    #Inputs: pad_x and pad_y, int, final dimensions of the outputted image.
    #Outputs: PIL Image object
    w,h = in_img.size
    left = math.floor((pad_x - w) / 2)
    top = math.floor((pad_y - h) / 2)
    out_img = Image.new(in_img.mode,(pad_x,pad_y),color = (255,255,255))
    out_img.paste(in_img,(left,top))
    return out_img

def padFolder(pad_x,pad_y,pad_dir = global_img_dir):
    #Pad all images in pad_dir to a uniform size
    #pad_x and pad_y are the dimensions to pad to, pad_dir is the directory of images to pad
    #Modifies saved images in place, no return.
    for a in os.listdir(pad_dir):
        if a != "Thumbs.db":
            new_img = pad(Image.open("{}/{}".format(pad_dir,a)),pad_x,pad_y)
            new_img.save("{}/{}".format(pad_dir,a))
            
def binarize(in_img,threshold = 200):
    #Binarize images for easier image processing
    #Takes a PIL Image as input and returns a PIL Image as output
    #Threshold is an integet from 0 to 255 controlling where pixels get set to black or white
    pixels = in_img.load()
    w,h = in_img.size
    for x in range(w):
        for y in range(h):
            avg = sum(pixels[x,y]) // 3
            if avg < threshold:
                pixels[x,y] = (0,0,0)
            else:
                pixels[x,y] = (255,255,255)
    return in_img

def binFolder(threshold = 200,bin_dir = global_img_dir):
    #Binarize all images in the inputted folder
    for a in os.listdir(bin_dir):
        if a != "Thumbs.db":
            new_img = binarize(Image.open("{}/{}".format(bin_dir,a)),threshold)
            new_img.save("{}/{}".format(bin_dir,a))

def convolve(in_img,kernel,mode = 0):
    #Do a convolution of a kernel across an image
    #The kernel is a square Python list of weights with odd dimensions (typically 3x3 or 5x5)
    #Some sample 3x3 kernels:
    '''
    Ridge detection
    [-1,-1,-1]
    [-1, 8,-1]
    [-1,-1,-1]
    Sharpen
    [ 0,-1, 0]
    [-1, 5,-1]
    [ 0,-1, 0]
    Gaussian blur
    [1/16, 2/16, 1/16]
    [2/16, 4/16, 2/16]
    [1/16, 2/16, 1/16]
    
    '''
    #Modes are 0 for sum, 1 for average, 2 for median
    #The output is a PIL Image object
    out_img = Image.new("RGB",in_img.size,color = (255,255,255))
    out_px = out_img.load()
    k_pad = (len(kernel) - 1)
    temp_img = pad(in_img,in_img.size[0] + 2*k_pad,in_img.size[1] + 2*k_pad)
    in_px = temp_img.load()
    #Once we have all out images loaded in, we can begin convoluting
    for x in range(out_img.size[0]):
        for y in range(out_img.size[1]):
            #The center of the kernel will be inside the dimensions of out_img, the padding will keep the kernel inside temp_img
            inter_list = []
            for a in range(len(kernel)):
                for b in range(len(kernel)):
                    #At each point here we need to do in_px[c,d] * kernel[a,b]
                    #This gets us the coordinates of the pixels within the kernel at point x,y
                    c = (x + (a - k_pad))
                    d = (y + (b - k_pad))
                    inter_list.append([e*kernel[a][b] for e in in_px[c,d]])
            #Now, what we do depends on mode
            if mode == 0:
                #Sum of all points, used for most convolutions
                out_img[x,y] = (sum([a[0] for a in inter_list]),sum([a[1] for a in inter_list]),sum([a[2] for a in inter_list]))
            elif mode == 1:
                #Same as before, but divide by kernel size
                out_img[x,y] = (int(sum([a[0] for a in inter_list]) / (len(kernel) ** 2)),int(sum([a[1] for a in inter_list]) / (len(kernel) ** 2)),
                             int(sum([a[2] for a in inter_list]) / (len(kernel) ** 2)))
            elif mode == 2:
                #Median filter
                out_px[x,y] = (sorted([a[0] for a in inter_list])[len(kernel) ** 2 // 2 + 1],sorted([a[1] for a in inter_list])[len(kernel) ** 2 // 2 + 1],
                             sorted([a[2] for a in inter_list])[len(kernel) ** 2 // 2 + 1])
    return out_img

def deNoise(in_img):
    #A common set of parameters for convolve
    return convolve(in_img,[[1,1,1],[1,1,1],[1,1,1]],2)
    

#__FUNCTIONS FOR GENERATING TYPEWRITTEN CHARACTER IMAGES__

def genChars(num_chars,char_set = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890",save_dir = global_img_dir):
    #Generate a random group of characters for testing purposes.
    #These will be saved with their identifier in the image file name, and used as training and validation samples.
    #Inputs are a string containing all characters in the input set, and an integer for how many you want.
    #__1: Create Word document containing random string of characters, separated by spaces__
    rand_string = " ".join([random.choice(char_set) for a in range(num_chars)])
    temp_file = docx.Document()
    temp_file.styles['Normal'].font.name = "Courier"
    temp_file.styles['Normal'].font.size = docx.shared.Pt(12)
    temp_file.styles['Normal'].paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.DOUBLE
    temp_file.add_paragraph().add_run(rand_string)
    temp_file.save("tempfile.docx")
    #__2: Create PDF from Word document, and save pages in that PDF as images__
    docx2pdf.convert("tempfile.docx","tempfile.pdf")
    temp_pdf = fitz.open("tempfile.pdf")
    chars = []
    for x,page in enumerate(temp_pdf):
        page_pxmp = page.get_pixmap(dpi = 100)
        page_img = Image.frombytes("RGB",[page_pxmp.width,page_pxmp.height],page_pxmp.samples)
        rows = chopList(page_img,0)
        tempcharlist = [chopList(a,1,[0,1]) for a in rows]
        for a in tempcharlist:
            chars += a
        #Theoretically, at this point, chars should be a list of relatively uniformly sized images of characters represented by rand_string
        
    #__3: Save letter images
    for i in range(len(rand_string.replace(" ",""))):
        chars[i].save("{}/{}_{}.png".format(save_dir,i,char_set.index(rand_string.replace(" ","")[i])))

def chopList(in_img,mode,offset = (0,0)):
    #Finds regions of in_img separated by lines of whitespace, used for segmenting text
    #At the moment is only equipped to segment by purely horizontal or vertical lines
    #in_img is a PIL Image object, and mode is 0 for row segmentation and 1 for column segmentation.
    #This is pretty oversensitive right now and could be made much better
    w,h = in_img.size
    pixels = in_img.load()
    boundary_list = []
    #__1: If mode 0 build boundary_list by rows, else build by columns__
    if mode == 0:
        #For each row, iterate through the row and determine if all pixels in the row are white
        for y in range(h):
            white_line = True
            for x in range(w):
                if pixels[x,y] != (255,255,255):
                    white_line = False
                    break
            boundary_list.append(white_line)
    elif mode == 1:
        #For each column, iterate through the column and determine if all pixels in the column are white
        for x in range(w):
            white_line = True
            for y in range(h):
                if pixels[x,y] != (255,255,255):
                    white_line = False
                    break
            boundary_list.append(white_line)
    #__2: Find regions by finding spots where boundary_list changes
    region_list = [a for a in range(len(boundary_list) - 1) if boundary_list[a] != boundary_list[a+1]]
    #__3: Make a list of image sub-regions
    if mode == 0:
        #If we're doing a row chop:
        segment_list = [in_img.crop((0,region_list[2*a] - offset[0],w,region_list[2*a+1]+offset[1])) for a in range(int(len(region_list) / 2))]
    elif mode == 1:
        #If we're doing a column chop:
        segment_list = [in_img.crop((region_list[2*a]-offset[0],0,region_list[2*a+1]+offset[1],h)) for a in range(int(len(region_list) / 2))]
    return segment_list

def findChars(in_img):
    #This function finds connected regions of black pixels surrounded by whitespace, and gets their bounding boxes
    #It should be run on binarized RGB PIL Images
    pixels = in_img.load()
    w,h = in_img.size
    #__1: Make a 2d list of all pixels, 1 if black, 0 if white
    pixel_list = [[0 if pixels[x,y] == (255,255,255) else 1 for x in range(w)] for y in range(h)]
    #Now, how do we get from here to finding objects?
    #An object is a region of connected pixels.
    #Pixels are connected if they are orthogonally or diagonally adjacent
    #__2: Find objects within pixel_list__
    objects = []
    for y in range(h):
        for x in range(w):
            if pixel_list[y][x] == 1:
                #If this is a black pixel, we need to check its neighbors
                neighborcoords = [(a,b) for a,b in zip(range(x-1,x+2),range(y-1,y+2))]
                #We're not actually trying to access pixel_list[b][a] at any point, so we don't need to sanitize this
                found = False
                for point in neighborcoords:
                    #For each neighboring coordinate:
                    for obj in objects:
                        if point in obj:
                            #If that point is in an existing object list, add this point to that object list
                            found = True
                            obj.append((x,y))
                if not found:
                    #If this is the first pixel of a new object,
                    objects.append([(x,y)])
    #__3: Return the center and bounding box of all objects__
    box_list = [[min([a[0] for a in b]),min([a[1] for a in b]),max([a[0] for a in b]),max([a[1] for a in b])] for b in objects]
    return [((int((a[2] - a[0])/2),int((a[1] - a[3])/2)),a) for a in box_list]

def boxDebug(in_img):
    #Debug function taking an image as input, and coloring the outlines of all characters found by findChars in red
    #__1: Run findChars to get the list of all objects__
    object_list = findChars(in_img)
    #__2: Make a new image to draw on__
    w,h = in_img.size
    out_img = Image.new(in_img.mode,(w,h),color = (255,255,255))
    out_img.paste(in_img)
    out_draw = ImageDraw.Draw(out_img)
    #__3: Draw the bounding boxes in__
    for box in object_list:
        out_draw.rectangle(box[1],outline = (255,0,0))
    out_img.save("rectDebug.img")
    
    

def imgToData(in_img,char_type,char_set = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890"):
    #This function takes an image, and turns it into numpy input/output pairs of neural network inputs
    #This should be run on binarized RGB PIL Images.
    #char_set is the possible characters in the dataset, used for transforming a symbol into ML output
    #The current version is working for letters and numbers, but we have some errors still on punctuation.
    #The return is a tuple of np arrays
    pixels = in_img.load()
    w,h = in_img.size
    #__1: Make a 1d list of all pixels, 1 if black, 0 if white__
    #Runs in reading order, left to right and top to bottom
    val_list = []
    for y in range(h):
        for x in range(w):
            if pixels[x,y] == (255,255,255):
                val_list.append(0)
            else:
                val_list.append(1)
    #__2: Make a 1d list of all 0s, except for 1 at the index of the correct answer
    c_index = char_set.index(char_type)
    e_out = [1 if a == c_index else 0 for a in range(len(char_set))]
    #__3: Convert lists to 2d 1 by n NP arrays and return
    return (np.array([val_list]).T,np.array([e_out]).T)

def imgToString(in_img,index,set_len):
    #This function takes an image, and turns it into a string that can be stored in a text file for fast access later
    #in_img is a binarized RGB PIL Image
    #index and set_len are used to generate the expected output string for the image
    pixels = in_img.load()
    w,h = in_img.size
    #__1: Make a 1d list of all pixels, 1 if black, 0 if white__
    #Runs in reading order, left to right and top to bottom
    val_list = []
    for y in range(h):
        for x in range(w):
            if pixels[x,y] == (255,255,255):
                val_list.append(0)
            else:
                val_list.append(1)
    #__2: Make a 1d list of all 0s, except for 1 at the index of the correct answer
    e_out = [1 if a == index else 0 for a in range(set_len)]
    #__3: Convert lists to strings, and return
    return "".join([str(a) for a in val_list]) + " " + "".join([str(a) for a in e_out])

def loadDataset(d_dir = global_img_dir,char_set = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890"):
    #This function runs imgToData on the whole d_path folder and returns a list of the resulting tuples
    out_data = []
    for a in os.listdir(d_dir):
        if a != "Thumbs.db":
            out_data.append(imgToData(Image.open("{}/{}".format(d_dir,a)),a.split(".")[0][-1],char_set))
    return out_data

def saveDataset(d_dir = global_img_dir,s_dir = "ArsData",filename = "set"):
    #This function runs imgToString on the whole d_dir folder, and writes the contents to one really big text file
    out_data = open("{}/{}.txt".format(s_dir,filename),"w")
    for a in os.listdir(d_dir):
        if a != "Thumbs.db":
            out_data.write(imgToString(Image.open("{}/{}".format(d_dir,a)),int(a.split(".")[0].split("_")[1])))
    out_data.close()

def checkSize(sz_path = global_img_dir):
    #This function checks the minimum and maximum dimensions of images in sz_path
    #This is used to determine how much to pad images when generating a new dataset
    i_x = []
    i_y = []
    for a in os.listdir(sz_path):
        if a != "Thumbs.db":
            new_img = Image.open("{}/{}".format(sz_path,a))
            i_x.append(new_img.size[0])
            i_y.append(new_img.size[1])
    print("X from {} to {}".format(min(i_x),max(i_x)))
    print("Y from {} to {}".format(min(i_y),max(i_y)))
    return max(i_x),max(i_y)

def clearData(d_dir = global_img_dir):
    #This function deletes all images in d_dir in preparation for creating a new, different set of data
    for a in os.listdir(d_dir):
        os.remove("{}/{}".format(d_dir,a))
        
def newDataset(char_set,size,d_dir = global_img_dir):
    #This function wraps several other functions to make creating a new dataset faster
    #__1: Clear the directory__
    print("Clearing")
    clearData(d_dir)
    #__2: Make a new dataset, chop, and save images to the directory__
    print("Generating")
    genChars(size,char_set,d_dir)
    #__3: Get the sizes of the images in the directory, and pad all to one bigger than the largest__
    x,y = checkSize(d_dir)
    print("Padding")
    padFolder(x+1,y+1,d_dir)
    #__4: Binarize the directory__
    print("Binarizing")
    binFolder(bin_dir = d_dir)
    #__5: Print and return the input and output sizes of the neural network required to process the dataset
    print("Insize: {}, outsize: {}".format((x+1)*(y+1),len(char_set)))
    return (x+1)*(y+1),len(char_set)

#__MAIN()__