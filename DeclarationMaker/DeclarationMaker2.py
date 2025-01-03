#DeclarationMaker2

#This document makes individual declarations

import docx
import os
import PyPDF2
import re
import random

#Section 1: Laying out a pleading document

def makeDeclaration(inDict,inDec,outFile,callInfo = {}):
    #Main function that makes a declaration using inDict and inDec, and saves it to outFile
    #inDict is a dictionary containing case data, inDec is the index string of a declaration in the decfolder, outFile is the filepath to write to
    #callInfo is used to give the blank pleading template a name and a footer
    #Function does not return anything, it does save a Word document to the path in the variable outfile
    #Template Index: 00 blank (must come with a title and footer in callInfo), 01 Dec Funds Rcvd, 02 Dec Lost Mail to DB, 03 Dec Lost Mail to POE,
    #04 Dec of Missing 1st, 05 Dec of Missing 2nd, 06 Dec of Frozen Interest, 07 Dec of Service, 08 Mil Status Dec, 09 Motion for Garn Costs (Fed POE),
    #10 Motion for Garn Costs (Missing 2nd Ans), 11 Order for Garn Costs, 12 Cert of Garn Costs, 13 Motion for Ext, 14 Order for Ext, 15 Dec re Ext, 16 Dec re Int,
    #17 Motion to Amend Name, 18 Dec for Amended Name, 19 WDS, 20 GR14 coversheet
    outdoc = docx.Document("programData/PleadingBlank.docx") #Open a file containing the right background and margins
    outdoc.styles['Normal'].font.name = "Courier" #Set to a constant-width font for the whole file for layout purposes
    outdoc.styles['Normal'].font.size = docx.shared.Pt(12) #Set to a constant size for layout purposes
    outdoc.styles['Normal'].paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.DOUBLE #Set to double-spaced by default
    #Read data from the declaration template to pass to other files later
    decFile = open(decByIndex(inDec),"r").read().splitlines()
    #Check that we have all the correct field names in inDict; if not then print an error message
    complete = True
    missingFields = []
    for item in ["Level","County","Division","Plaintiff","Defendant","case#","Packet"]:
        if item not in inDict.keys():
            missingFields.append(item)
            complete = False
    if not complete:
        print("Error: Missing",missingFields)
        return
    else:
        #If complete: Create docx document using data from inDict and the relevant helper functions
        if callInfo == {}:
            headerName = decFile[0].strip()
            footerName = decFile[1].strip().replace("{ACTION}","")
        else:
            headerName = callInfo["Header"]
            footerName = callInfo["Footer"]
        header(inDict,headerName,document = outdoc,action = ("{ACTION}" in decFile[1]))
        body(inDict,decFile[2:],document = outdoc, doctitle = footerName)
        footer(inDict,footerName.strip().replace("{ACTION}",""),document = outdoc)
        #Once created, save file
        outdoc.save(outFile)
        #print("Saved to {}".format(outFile))
        return

def header(inDict,docTitle,document = docx.Document(),debugOut = False,action = False):
    #Lay out a header and caption, and add them to the inputted document
    #inDict is the dictionary containing case data
    #docTitle is the title of the pleading
    #document is the docx.Document object we're writing to
    #debugOut prints the result of the header to the console for observation
    #action adds additional language to the title if true
    #Returns: the function modifies the Document object in place, and returns the debug list in case anyone needs it
    debugOutlist = []
    for x in range (4): #Space out the header to leave space for the court to stamp, they get mad if this isn't big enough
        document.add_paragraph()
    #Create a court header from the information in inDict
    headerLine = "IN THE {} COURT OF THE STATE OF WASHINGTON\nIN AND FOR {} COUNTY".format(inDict["Level"].upper(),inDict["County"].upper())
    if inDict["Division"] != "": headerLine += ", {} DIVISION".format(inDict["Division"].upper())
    debugOutlist.append(headerLine)
    #Write the header to a new center-aligned paragraph
    header = document.add_paragraph()
    header.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    header.add_run(headerLine)
    #Caption (parties box)
    #At the moment we're using a fixed-width font to make the header box uniform in width
    partytypes = ["Plaintiff","Defendant(s)","Garnishee"]
    partyboxwidth = 34 #Enough width for the party box to the right to contain all party information
    docwidth = 60 #The approximate width of a Word document in 12pt Courier characters, if font/size changed this would need to as well
    partyboxlines = []
    #Sets up party header; if there are more parties the box is taller
    if "Garnishee" in inDict.keys(): boxdepth = 5
    else: boxdepth = 4
    for x in range (boxdepth):
        if x in [0,boxdepth-1]: 
            partyboxlines.append("-"*partyboxwidth) #Top and bottom of the box
        else:
            #This sets up the program to write a blank line after the party type and makes single-line parties work properly
            currentparty = inDict[partytypes[x-1].replace("(s)","")] + "| |" + partytypes[x-1] + "| "
            templines = currentparty.split("|")
            for y in range(len(templines)):
                if y == len(templines) - 2: #Partytype is right-aligned within the party box
                    partyboxlines.append("{0: >{1}} |".format(templines[y],partyboxwidth - 2))
                elif y == len(templines) - 4: #Comma at the end of the party
                    partyboxlines.append("{0: <{1}} |".format(templines[y] + ",",partyboxwidth - 2))
                else: #Otherwise just add spaced content
                    partyboxlines.append("{0: <{1}} |".format(templines[y],partyboxwidth - 2))
        if x == 1:
            #After the plaintiff add "vs." centrally aligned in the partybox
            partyboxlines.append("{0: ^{1}} |".format("vs.",partyboxwidth - 2))
            partyboxlines.append("{0: <{1}} |".format("",partyboxwidth - 2))
    #We now have the left side of the header containing the parties, right side contains case number, pleading title, and packet number
    #Currently no edge case handling for really long pleading titles, so be careful
    #Splits the title into a set of fragments that fit at the end of each partyboxline based on document width
    titlewords = docTitle.split(" ")
    titlelist = []
    titlebuffer = titlewords[0]
    for x in range (1,len(titlewords)):
        if len(titlebuffer) + len(titlewords[x]) + 3 >= docwidth - partyboxwidth:
            titlelist.append(titlebuffer)
            titlebuffer = titlewords[x]
        else:
            titlebuffer += " " + titlewords[x]
    titlelist.append(titlebuffer)
    if action:
        titlelist.append("")
        titlelist.append("(CLERK'S ACTION REQUIRED)")
    #The caseno and packet lines come from inDict and should always fit on one line
    #Add the stuff on the right of the party box to the party box
    for x in range(len(partyboxlines)):
        if x == 1: partyboxlines[x] += "  NO. {}".format(inDict["case#"])
        elif x in range(3,len(titlelist) + 3): partyboxlines[x] += "  {}".format(titlelist[x-3])
        elif x == len(partyboxlines) - 2: partyboxlines[x] += "  (P{}/NXS)".format(inDict["Packet"])
        partyboxlines[x] = partyboxlines[x].replace(",,",", ")
        partyboxlines[x] += "\n"
    #Write the party box to the document in a single-spaced paragraph to save on height
    partyboxpar = document.add_paragraph()
    partyboxpar.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    partyboxpar.paragraph_format.space_after = docx.shared.Pt(0)
    for x in range(len(partyboxlines)):
        if x in range (3,len(titlelist) + 3):
            partyboxpar.add_run(partyboxlines[x].split("|")[0] + "|")
            partyboxpar.add_run(partyboxlines[x].split("|")[1]).bold = True
        else:
            partyboxpar.add_run(partyboxlines[x])
        debugOutlist.append(partyboxlines[x])
    #If in debug mode, print the whole party box
    if debugOut:
        for x in range(len(debugOutlist)):
            if x == 0: print(debugOutlist[x])
            else: print(debugOutlist[x],end="") #We need end="" here because the partybox has newlines at the end of each line
    return debugOutlist
    

def body(inDict,inTemplate,document = docx.Document(),debugOut = False,doctitle = ""):
    #Write the pleading body to the document
    #inDict is the case data to fill in data fields in the template
    #inTemplate is lines 2: of the template txt file (lines 0 and 1 are title info for the header function)
    #document is the Word document to modify
    #{} contain flags data fields, or special instructions. I don't have a comprehensive list at present, more are added as they're needed
    debugList = []
    #Formatting flags to allow sections of the document to be reformatted
    spflag = False
    rflag = False
    indflag = True
    blockwidth = 55 #Used for spacing out certain kinds of primitive tables in the document
    for line in inTemplate:
        templine = line #New string, because it can be modified
        #If the whole line is one of these commands the line is not written to the document, either a flag changes or something else is printed
        if templine.strip() == "{SsBlock}": spflag = not spflag
        elif templine.strip() == "{Unindblock}": indflag = not indflag
        elif templine.strip() == "{Rightblock}": rflag = not rflag
        elif templine.strip() == "{Break}": document.add_paragraph().add_run().add_break(break_type = docx.enum.text.WD_BREAK.PAGE)
        elif templine.strip() == "{AddrBlock}":
            addrpar = document.add_paragraph()
            addrpar.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
            addrpar.paragraph_format.left_indent = docx.shared.Inches(0.5)
            addrpar.add_run("{}\n{}\n{}".format(inDict["DefS"].strip(),inDict["AddrD"].split(", ")[0].strip(),", ".join(inDict["AddrD"].split(", ")[1:]).strip()))
        elif templine.strip() in ["{Mesigblock}","{Mikesigblock}","{Judgesigblock}"]:
            sigblock(inDict,templine.strip(),document = document)
        #If the whole line is not one of those special commands:
        else:
            #Add a paragraph to the document
            nextpar = document.add_paragraph()
            boldflag = False
            #If the line contains any parameters that are inDict keys contained in {} replace the field with the data
            if "{" in templine:
                for parameter in inDict.keys():
                    parcheck = "{" + parameter + "}"
                    if parcheck in templine and inDict[parameter] != "":
                        templine = templine.replace(parcheck,parparse(inDict,parameter))
            #Other specific formatting and data {} patterns.
            if "{DOC}" in templine:
                templine = templine.replace("{DOC}",doctitle)
            if "{SpaceLine}" in templine:
                templine = "  " + templine.split("|")[1] + "{0: >{1}}".format(templine.split("|")[2].strip(),blockwidth - len(templine.split("|")[1]))
            if "{Bold}" in templine:
                templine = templine.replace("{Bold}","")
                boldflag = True
            if "{DIVISION?}" in templine and inDict["Division"]: templine = templine.replace("{DIVISION?}",", {} Division".format(inDict["Division"][0].upper() + inDict["Division"][1:].lower()))
            else: templine = templine.replace("{DIVISION?}","")
            if "{Center}" in templine:
                templine = templine.replace("{Center}","")
                nextpar.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            debugList.append(templine)
            #Change paragraph formatting and write templine to document
            if spflag:
                nextpar.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
            if indflag:
                nextpar.paragraph_format.first_line_indent = docx.shared.Inches(0.5)            
            if rflag:
                nextpar.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                if boldflag: nextpar.add_run(templine.split("|")[0]).bold = True
                else: nextpar.add_run(templine.split("|")[0])
                nextpar.add_run(templine.split("|")[1]).font.color.rgb = docx.shared.RGBColor(0xff,0xff,0xff)
            else:
                nextpar.add_run(templine).bold = boldflag            
            #Debug check for leftover unhandled {} tags
            errors = re.findall("\{.*?\}",templine)
            if len(errors) > 0:
                print("Missed fields in " + inDict["case#"] + ": ",errors)
            boldflag = False
    #if debug, print debug info to save on opening Word docs
    if debugOut: [print(entry,end="") for entry in debugList]
    return debugList

def sigblock(inDict,sigtype,document = docx.Document(),debugOut = False):
    #Format and add to document signature blocks for judge, attorney, or assistant
    blockwidth = 40
    debuglist = []
    #This segment is the date line
    dateblock = document.add_paragraph()
    dateblock.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    #Judges sign things whenever they get around to them, and so need a blank line instead of the date we printed the document
    if sigtype == "{Judgesigblock}": dateblock.add_run("SO ORDERED this ______ day of _________________, ________")
    else: dateblock.add_run("Dated at Redmond, Washington this " + parparse(inDict,"DtDc") + ".")
    #This segment is the signature section
    sigblock = document.add_paragraph()
    sigblock.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
    sigblock.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
    runlist = []
    #The actual signature lines
    if sigtype == "{Mesigblock}":
        blockwidth = 21
        runlist += ["","____________________",parparse(inDict,"Me"),"Legal Case Manager"]
    elif sigtype == "{Mikesigblock}":
        blockwidth = 36
        #Mike has more space above his signature line because he signs things really big and his stamp is really big
        runlist += ["Presented by:","   {B}O'MEARA LAW OFFICE P.S.","   {B}AT MERCHANTS CREDIT CORPORATION","","","","   By:_________________________","     Michael S. O'Meara, WSBA #41502","     Attorney for Plaintiff"]
    elif sigtype == "{Judgesigblock}":
        blockwidth = 24
        #Judges have more space above their signature line...just in case?
        runlist += ["","","","________________________","{B}Judge/Court Commissioner"]
    for line in runlist:
        if "{B}" in line: boldflag = True
        else: boldflag = False
        templine = "{0: <{1}}".format(line.replace("{B}",""),blockwidth)
        debuglist.append(templine + "|")
        #This part is important for right-aligning the signature block correctly
        #If you just pad to the end with spaces Word will let the spaces go offscreen without changing anything, so we need a character at the end of the line
        #But we don't want to see a line of extra pipes, so we print them in white
        sigblock.add_run(templine).bold = boldflag
        sigblock.add_run("|\n").font.color.rgb = docx.shared.RGBColor(0xff,0xff,0xff)
    if debugOut:
        for line in debuglist:
            print(line)

def footer(inDict,docTitle,document = docx.Document(),debugOut = False):
    #I don't understand exactly what most of this does, but the end result is a footer with the document title and a page number
    quickrun = "L{} P{} {} PAGE ".format(inDict["Legal"],inDict["Packet"],docTitle)
    fldChar1 = docx.oxml.OxmlElement("w:fldChar")
    fldChar1.set(docx.oxml.ns.qn("w:fldCharType"),"begin")
    instrText = docx.oxml.OxmlElement("w:instrText")
    instrText.set(docx.oxml.ns.qn("xml:space"),"preserve")
    instrText.text = "PAGE"
    fldChar2 = docx.oxml.OxmlElement("w:fldChar")
    fldChar2.set(docx.oxml.ns.qn("w:fldCharType"),"end")
    foot = document.sections[0].footer.paragraphs[0].add_run(quickrun)
    foot._r.append(fldChar1)
    foot._r.append(instrText)
    foot._r.append(fldChar2)    



#Section 2: Helper Functions

def parparse(inDict,parameter,debug = False):
    #This formats data from our inDict fields to the correct longform for use in the documents
    if debug: print("Attempting to parse:",parameter)
    par = inDict[parameter]
    #These are helper lists for converting dates
    dateapp = ["th","st","nd","rd","th","th","th","th","th","th"]
    months = ["","January","February","March","April","May","June","July","August","September","October","November","December"]
    if parameter[:2] == "Dt" and par != "":
        if str(int(par[2:4])) in ["11","12","13"]:
            #Added to fix "11st","12nd","13rd"
            pardate = par[2:4] + "th"
        else:
            pardate = str(int(par[2:4])) + dateapp[int(par[2:4])%10]
        parmo = months[int(par[:2])]
        paryear = "20" + par[4:] #If this program is still in use in 87 years I will be very, very surprised
    if parameter == "DtDc" and par != "": #One specific case in the signature line wants "18th day of September, 2023"
        return pardate + " day of " + parmo + ", " + paryear
    elif parameter[:2] == "Dt" and par != "": #Most date cases want "September 18th, 2023"
        return parmo + " " + pardate + ", " + paryear
    elif parameter in ["Level","County","Division"]: #These are stored in the dictionary in all-caps
        return par[0].upper() + par[1:].lower()
    else: #If it's not a parameter that needs changing just return it as is
        return par
        
def decByIndex(index):
    #Helper function to convert from two-numeral index string to full relative path
    decsList = os.listdir("programData/PleadingTemplates")
    for file in decsList:
        if file[:2] == index[:2]:
            return "programData/PleadingTemplates/" + file

def makeLabels(data1,data2,outname = "MailingLabels.docx"):
    #Makes mailing labels from a data1 dict (outputted by newParser) and a data2 list (outputted by extractRAg)
    #This uses a blank document from programData that has the correct margins to print mailing labels
    #Table is 3 wide, has 10 rows per page, and we need five labels per case
    document = docx.Document("programData/LabelsBlank.docx")
    table = document.tables[0] #Get the table from the document
    print(len(data1) * 5 // 3 + 1) #Print the number of pages the output should be for debugging purposes
    for x in range((len(data1) * 5) // 3 - 9):
        #Add additional rows to the table until the table contains enough cells for five labels per case
        newrow = table.add_row()
        newrow.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.EXACTLY
        newrow.height = docx.shared.Inches(1)
    for x in range(len(data1)):
        #For each case, add labels to the table
        #Pull info from data1, using getData for quick error handling
        Caseno = getData(data1[x],"case#")
        Gname = getData(data1[x],"Garnishee")
        Dname = getData(data1[x],"DefS")
        #Format addresses with street, csz on separate lines
        tempaddrG = getData(data1[x],"AddrG").split(",")
        Gaddr1 = tempaddrG[0]
        Gaddr2 = ",".join(tempaddrG[1:])
        tempaddrD = getData(data1[x],"AddrD").split(",")
        Daddr1 = tempaddrD[0]
        Daddr2 = ",".join(tempaddrD[1:])
        #If we have a Gname we can get the registered agent from data2, otherwise we need to recheck manually
        if "ERROR" not in Gname:
            RagRaw = data2[x][fuzzyMatch(data2[x],Gname)]
            RAg = "{}".format(RagRaw[1].split("  ")[-1])
            if ":" not in RAg: RAg = "ATTN: " + RAg #Account for people writing "ATTN:" into the registered agent line
        else:
            RAg = "ERROR: REVIEW"
        Grun = "{}\n{}\n{}\n{}\n{}".format(Caseno,Gname,RAg,Gaddr1,Gaddr2)
        Drun = "{}\n{}\n{}\n{}".format(Caseno,Dname,Daddr1,Daddr2)
        #Write three garnishee mailing labels and two defendant mailing labels to the table
        writelist = [Grun,Grun,Grun,Drun,Drun]
        for y in range (x*5, (x+1)*5):
            crow = y//3
            ccolumn = 2*(y%3)
            run = writelist[y-(x*5)]
            table.cell(crow,ccolumn).add_paragraph(run)
    document.save(outname)


def fuzzyMatch(inList,pattern):
    #Checking for matching words between two strings to account for formatting changes
    #Not really generalizable from this in application, I don't think
    testPattern = pattern.split(" ")
    outIndex = [0] * len(inList)
    for word in testPattern:
        for x in range(len(outIndex)):
            if word in inList[x][0]:
                outIndex[x] += 1
    maxIndex = 0
    for x in range(len(outIndex)):
        if outIndex[x] > outIndex[maxIndex]:
            maxIndex = x
    return x

def getData(inDict,key):
    #Helper function to handle key errors without stopping the run of the program
    try: return inDict[key]
    except: return "KEY ERROR {}".format(key)


#Section 3: Other Projects

def envelopeMaker(addrList):
    #In theory a function to make printable envelopes, but I have yet to make it actually print envelopes
    #Landscape page, 2" left margin, 2" right margin, 
    #9.5" wide, 4.125" tall
    outDoc = docx.Document()
    current_section = outDoc.sections[-1]
    current_section.page_width = docx.shared.Inches(9.5)
    current_section.page_height = docx.shared.Inches(4.125)
    current_section.top_margin = docx.shared.Inches(2.25)
    current_section.bottom_margin = docx.shared.Inches(0.5)
    current_section.left_margin = docx.shared.Inches(3.25)
    current_section.right_margin = docx.shared.Inches(3.25)
    for address in addrList:
        contentpar = outDoc.add_paragraph()
        #contentpar.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        for entry in address.split("|"):
            contentpar.add_run("{}\n".format(entry))
        contentpar.add_run().add_break(break_type = docx.enum.text.WD_BREAK.PAGE)
    outDoc.save("testEnvelope.docx")

def scrambleset():
    #Helper function for a tangentially related project, creates a set of 10,080 random characters for an OCR character identification experiment
    scrambledoc = docx.Document()
    scrambleindex = open("ocrtest/scrambleset/scrambleIndex1.txt","wb")
    scrambleindex.write(u"\uFEFF".encode("UTF-8"))
    chrset = [x for x in "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890-_!#$%&*()[]|,.?/\\\'\""]
    randchars = ""
    for x in range(10080):
        randcharchoice = random.choice(chrset) + " "
        print(randcharchoice)
        randchars += randcharchoice
        scrambleindex.write(randcharchoice.encode("UTF-8"))
    randpar = scrambledoc.add_paragraph()
    scrambledoc.styles['Normal'].font.name = "Courier"
    scrambledoc.styles['Normal'].font.size = docx.shared.Pt(12)
    scrambledoc.styles['Normal'].paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.DOUBLE
    randpar.add_run(randchars)
    #scrambleindex.write(randchars)
    scrambleindex.close()
    scrambledoc.save("Scrambledoc1.docx")
    
    