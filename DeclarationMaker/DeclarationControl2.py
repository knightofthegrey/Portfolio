#Revised Declaration Maker Tool
#Adjusted and cleaned up tools from v1
#Py files in Project: 
#DeclarationControl2 is the main file
#KeyInput2 uses the Python keyboard module to skim quickly through PAMAR and collect data
#DataProcessor2 takes text files outputted from PAMAR and turns them into dictionaries.
#DeclarationMaker2 takes dictionaries or text files outputted from DataProcessor and turns them into Word documents or PDFs
#PdfHandler2 concatenates and renames PDFs
#OCRrun uses OCR to split PDFs
#Folder Structure:
#inputData contains .txt and .pdf input files
#outputData contains documents that need to be printed or filed, typically collected in Batch folders of case folders
#programData contains templates that the program uses to make things

import PdfHandler2
import KeyInput2
import DataProcessor2
import DeclarationMaker2
import ConstantWidthTable
import docx
import docx2pdf
import re
import os
import PyPDF2
import shutil
import OCRrun
import subprocess
import Utilities
from pathlib import PureWindowsPath
from time import time

#Paths for various purposes
filedpath = "//diskstation/Legal Team/TEAM MEMBER DIRECTORIES/NXS/Filed Documents"
abspath = r"\\diskstation\Legal Team\TEAM MEMBER DIRECTORIES\NXS\Python Document Processing\Declaration Maker"
oldpath = "//diskstation/Legal Team/TEAM MEMBER DIRECTORIES/NXS/Old Inputs"
packetpath = r"\\diskstation\Global_Share\Scan_users\NXS"
docspath = r"C:\Users\nathanielslivke\Documents"

def makeQuickFolders(inData,batchnum = 0):
    #Makes folders for all cases in inData in the indicated batch
    if "Batch{}".format(batchnum) not in os.listdir("outputData"):
        os.mkdir("outputData/Batch{}".format(batchnum))
    for case in inData:
        if case["County"] not in ["KETCHIKAN","LANE"]:
            os.mkdir("outputData/Batch{}/{}".format(batchnum,case["FolderName"]))

def smartDecs(inData,now,mildate,batchnum = 0,me = "Nathaniel Slivka"):
    #This function checks the contents of any folders in outputData/Batch and uses those to determine which declarations are necessary for a regular order
    #We haven't needed a special order for a while, so this function isn't set up to do them
    #Files which could be in a folder for a regular order: 00 motion, 01 order, 02 2nd A, 03 1st A, 04 decS, 05 PoE missing, 06 DB missing, 07 FrInt, 08 Costs, 09 DecMil
    #00 motion, 01 order, either 03 1st A or 03 Dec of Missing 1st, 04 decS, 08 Costs, 09 decMil are in all orders
    #02 is 2nd Answer for regular orders or Dec Funds Rcvd for special orders, but there is always an 02
    #05 and 06 are missing mailings declarations for if 04 doesn't have all mail returns
    #07 references Proclamation 20-49 of 2020 and doesn't apply to judgments from after mid-2020
    qp = "outputData/Batch{}".format(batchnum)
    folders = os.listdir(qp)
    errorCleanup(mode = 1,batchnum = batchnum)
    for item in inData:
        item["Me"] = me
        item["DtMil"] = mildate
        item["DtDc"] = now
        item["DtDcA"] = now
        lp = qp + "/" + item["FolderName"]
        if item["FolderName"] in os.listdir(qp):
            folderContents = os.listdir(lp)
            if "03.pdf" not in folderContents:
                DeclarationMaker2.makeDeclaration(item,"04",lp + "/03 " + item["case#"] + " Missing 1st.docx")
            if "04_1.pdf" not in folderContents:
                DeclarationMaker2.makeDeclaration(item,"03",lp + "/05_0 "+item["case#"] + " Missing PoE.docx")
                mailnos = open(lp + "/mailnos.txt","a")
                mailnos.write(item["MailnoG"] + "\n")
                mailnos.close
            if "04_2.pdf" not in folderContents:
                DeclarationMaker2.makeDeclaration(item,"02",lp + "/06_0 " + item["case#"] + " Missing DB.docx")
                mailnos = open(lp + "/mailnos.txt","a")
                mailnos.write(item["MailnoD"] + "\n")
                mailnos.close
            if item["IntA"] != "":
                DeclarationMaker2.makeDeclaration(item,"06",lp + "/07_0 " + item["case#"] + " Fr Interest.docx")
            DeclarationMaker2.makeDeclaration(item,"07",lp + "/04_0 " + item["case#"] + " Dec of Serv.docx")
            item["DtDc"] = item["DtMil"]
            DeclarationMaker2.makeDeclaration(item,"08",lp + "/09_0 " + item["case#"] + " Dec Mil Status.docx")
            print("Case done:",item["case#"])

def extWrapper(inData,now,postdate = "",batchnum = 0,me = "Nathaniel Slivka",batched = True):
    #Make extension declarations and then set them up to print for signature all at the same time!
    extDecs(inData,now,postdate = postdate,batchnum = batchnum,me = me,batched = batched)
    PdfHandler2.extPrintBatch(batchnum)

def extDecs(inData,now,postdate = "",batchnum = 0,me = "Nathaniel Slivka",batched = True):
    #Make all necessary declarations for an extension for all folders in a batch
    #This one isn't set up to do non-batched, the extra argument is redundant
    qp = "outputData/Batch{}".format(batchnum)
    folders = os.listdir(qp)
    names = ["00 {} Motion.docx","01 {} Order.docx","02 {} Dec re Ext.docx","03 {} Dec re Int.docx","04 {} WDS.docx"]
    for item in inData:
        if item["FolderName"] in folders:
            item["Me"] = me
            if postdate: item["DtDc"] = postdate
            else: item["DtDc"] = now
            item["DtDcA"] = item["DtDc"]
            item["DtRun"] = now
            if "Washington" in item["Plaintiff"]: item["Plaintiff"] = "MERCHANTS CREDIT CORPORATION"
            lp = "{}/{}".format(qp,item["FolderName"])
            if batched:
                for x in range (13,17):
                    DeclarationMaker2.makeDeclaration(item,str(x),"{}/{}".format(lp,names[x-13].format(item["case#"])))
                DeclarationMaker2.makeDeclaration(item,"19","{}/{}".format(lp,names[4].format(item["case#"])))
            if "JOA" in item["f/desc"]: print("DOUBLE CHECK GARNS FOR {}".format(item["Legal"]))

def quickDecs(inData,now,mildate,dec,outname,batchnum = 0,me = "Nathaniel Slivka",batched = True):
    #Make one declaration for each folder in the batch
    qp = "outputData/Batch{}".format(batchnum)
    folders = os.listdir(qp)
    for item in inData:
        item["Me"] = me
        item["DtMil"] = mildate
        item["DtDc"] = now
        item["DtDcA"] = now
        lp = qp + "/" + item["FolderName"]
        if batched:
            DeclarationMaker2.makeDeclaration(item,dec,lp + "/" + outname.split("|")[0] + item["case#"] + outname.split("|")[1] + ".docx")
        else:
            DeclarationMaker2.makeDeclaration(item,dec,"outputData/{}{}{}.docx".format(outname.split("|")[0],item["case#"],outname.split("|")[1]))

def workaroundLnums(inFile,outFile):
    #Write all legal numbers from infile to outfile
    inFilepath = "inputData/{}".format(inFile)
    outFilepath = "outputData/{}".format(outFile)
    inputData = open(inFilepath,"r").read().splitlines()
    outFile = open(outFilepath,"w")
    outputList = []
    for line in inputData:
        if re.search("^\d+ +L\d+ P\d+$",line) and line not in outputList:
            outputList.append(line)
            outFile.write("{} {}\n".format(line.split(" ")[-2][1:],line.split(" ")[-1][1:]))

def quickprintall():
    #Merge all PDFs in the main folder for printing
    merger = PyPDF2.PdfMerger()
    for file in os.listdir():
        if "." in file and file.split(".")[-1] == "docx":
            docx2pdf.convert(file,file.split(".")[0] + ".pdf")
            merger.append(file.split(".")[0] + ".pdf")
    merger.write("mergefile.pdf")

def errorCleanup(mode = 0,batchnum = 0):
    #Fixes naming errors in batches
    qp = "outputData/Batch{}".format(batchnum)
    if mode == 0:
        #Obsolete, corrected for an older version of a splitter
        for subfolder in os.listdir(qp):
            for file in os.listdir("{}/{}".format(qp,subfolder)):
                if "SIGNED" in file:
                    os.remove("{}/{}/{}".format(qp,subfolder,file))
    elif mode == 1:
        #Necessary due to weird quirks of how the capture process works
        #"Shift" is the key that writes the next legal number to the keyboard when using KeyInput to request documents,
        #so I can't type _ when naming files coming from PAMAR, but the program expects 04 1 and 04 2 to be 04_1 and 04_2
        for subfolder in os.listdir(qp):
            for file in os.listdir("{}/{}".format(qp,subfolder)):
                if file in ["04 1.pdf","04 2.pdf"]:
                    os.rename("{}/{}/{}".format(qp,subfolder,file),"{}/{}/{}".format(qp,subfolder,file.replace(" ","_")))
    elif mode == 2:
        #If something was very wrong with the whole declaration maker process clear all non-pdfs from the batch and start over
        for subfolder in os.listdir(qp):
            for file in os.listdir("{}/{}".format(qp,subfolder)):
                if ".pdf" not in file:
                    os.remove("{}/{}/{}".format(qp,subfolder,file))
    elif mode == 3:
        #If I accudentally ran the declaration maker too early remove extraneous missing mailing declarations
        for subfolder in os.listdir(qp):
            subpath = "{}/{}".format(qp,subfolder)
            if "04_1.pdf" in os.listdir(subpath):
                for file in os.listdir(subpath):
                    if "05_0" in file:
                        os.remove("{}/{}".format(subpath,file))
            if "04_2.pdf" in os.listdir(subpath):
                for file in os.listdir(subpath):
                    if "06_0" in file:
                        os.remove("{}/{}".format(subpath,file))
    elif mode == 5:
        #Correcting a bug that gave duplicate WDS to some extensions
        for subfolder in os.listdir(qp):
            subpath = "{}/{}".format(qp,subfolder)
            for file in os.listdir(subpath):
                if "WDS.pdf" in file:
                    os.remove("{}/{}".format(subpath,file))

def inputFinder(batchnum = 0):
    #Old file to try and identify input txts from before the current naming schemes
    for file in os.listdir("inputData"):
        if ".txt" in file:
            try:
                data = DataProcessor2.newParser(file)
                if data[1]["FolderName"] in os.listdir("outputData/Batch{}".format(batchnum)) and abs(len(data) - len(os.listdir("outputData/Batch{}".format(batchnum)))) <= 1:
                    return data
            except:
                pass
    return "No match"

def garnPackets(inFolder):
    #Splits new garnishment documents in preparation for setting them up to print
    #Expecting: 1st Answers.pdf, Applications.pdf, Exemptions.pdf, Writs.pdf
    print("Run1")
    answerIndex = OCRrun.fasterSplit("{}/1st Answers.pdf".format(inFolder),4)
    print("Run2")
    appIndex = OCRrun.fasterSplit("{}/Applications.pdf".format(inFolder),2)
    print("Run3")
    exIndex = OCRrun.fasterSplit("{}/Exemptions.pdf".format(inFolder),2)
    print("Run4")
    writsIndex = OCRrun.fasterSplit("{}/Writs.pdf".format(inFolder),5)
    caseIndex = {}
    debugList = {}
    #Build casenoIndex from writs:
    for file in writsIndex:
        caseIndex[file[0]] = [file[1]]
    #Append answers, apps, exemptions in that order
    for subfile in [answerIndex,appIndex,exIndex]:
        for file in subfile:
            if file[0] in caseIndex.keys():
                caseIndex[file[0]].append(file[1])
            else:
                if file[0] in debugList.keys(): debugList[file[0]].append(file[1])
                else: debugList[file[0]] = [file[1]]
    #Check output:
    print("Found {} cases".format(len(caseIndex)))
    for caseno in caseIndex.keys():
        print("Case {} contains {} files".format(caseno,len(caseIndex[caseno])))
    print("\n==MISMATCHES==\n")
    for key in debugList.keys():
        print("Mismatched case: {} contains {} files".format(key,len(debugList[key])))

def dictAppend(inDict,key,value):
    #This is now in Utilities, but this file hasn't entirely been converted over
    outDict = inDict
    if key in outDict.keys(): outDict[key].append(value)
    else: outDict[key] = [value]
    return outDict

def garnCountyMerger():
    #Obsolete, new version is in PdfHandler
    path = "inputData/New Garn Processing/Final/FileFolder"
    tempData = os.listdir(path)
    countyDict = {}
    for file in tempData:
        fileData = OCRrun.getInfo("{}/{}".format(path,file),0,3)
        if fileData["Level"] == "SUPERIOR": countyDict = dictAppend(countyDict,"SUPERIOR",file)
        else: countyDict = dictAppend(countyDict,fileData["County"],file)
    groupedPrint = PyPDF2.PdfMerger()
    for key in countyDict.keys():
        if key in ["SUPERIOR","KING","BENTON","YAKIMA"]:
            for case in countyDict[key]:
                os.rename("{}/{}".format(path,case),"{}/{} {}".format(path,key,case))
        else:
            for case in countyDict[key]:
                groupedPrint.append("{}/{}".format(path,case))
    groupedPrint.write("{}/County Print Group.pdf".format(path))

def diagnose(batchnum = 0):
    #Finds cases with mail numbers that need to be requested from USPS
    path = "outputData/Batch{}".format(batchnum)
    for file in os.listdir(path):
        subdata = os.listdir("{}/{}".format(path,file))
        if "mailnos.txt" in subdata:
            print(file)
            print(open("{}/{}/{}".format(path,file,"mailnos.txt"),"r").read().splitlines())

def openBatchForSig(batchnum = 0,patterns = ["\d_0"],batchnames = []):
    #Opens PDFs in a given batch so they can be signed
    #Mostly obsolete, new version merges them to avoid all the Acrobat loading time
    abspath = r"\\diskstation\Legal Team\TEAM MEMBER DIRECTORIES\NXS\Python Document Processing\Declaration Maker\outputData\Batch{}".format(batchnum)
    path = "outputData/Batch{}".format(batchnum)
    acropath = r"C:\Program Files\Adobe\Acrobat DC\Acrobat\Acrobat.exe"
    openfiles = 0
    for file in os.listdir(path):
        if len(batchnames) == 0 or file in batchnames:
            subpath = "{}/{}".format(path,file)
            subdata = os.listdir(subpath)
            for subfile in subdata:
                for pattern in patterns:
                    if re.search(pattern,subfile) and ".pdf" in subfile:
                        if openfiles >= 6:
                            input("Press enter to continue:")
                            openfiles = 0
                        print("Opening {}".format(subfile))
                        tempabspath = r"{}\{}\{}".format(abspath,file,subfile)
                        temppath = "{}/{}".format(subpath,subfile)
                        #subprocess.Popen(["Acrobat.exe",tempabspath])
                        os.startfile(tempabspath)
                        openfiles += 1

def openWordForFix(batchnum = 0,patterns = ["\d_0"],batchnames = []):
    #Open Word documents with a given filename pattern in the indicated batch if we need to correct lots of them manually for some reason
    abspath = r"\\diskstation\Legal Team\TEAM MEMBER DIRECTORIES\NXS\Python Document Processing\Declaration Maker\outputData\Batch{}".format(batchnum)
    path = "outputData/Batch{}".format(batchnum)
    openfiles = 0
    for folder in os.listdir(path):
        if len(batchnames) == 0 or folder in batchnames:
            subpath = "{}/{}".format(path,folder)
            subfolder = os.listdir(subpath)
            for subfile in subfolder:
                for pattern in patterns:
                    if re.search(pattern,subfile) and ".docx" in subfile:
                        if openfiles >= 6:
                            input("Press enter to continue:")
                            openfiles = 0
                        print("Opening {}".format(subfile))
                        tempabspath = r"{}\{}\{}".format(abspath,folder,subfile)
                        temppath = "{}/{}".format(subpath,subfile)
                        os.startfile(tempabspath)
                        openfiles += 1

def sortedDesknums(inData):
    #Writes desk numbers to file to help with check requests
    desknums = {}
    for y in inData:
        desknums = dictAppend(desknums,y["County"],y["dsk"])
    outFile = open("DesknumOut.txt","w")
    for county in desknums.keys():
        outFile.write("\n{}\n".format(county))
        for dsk in sorted(desknums[county]):
            outFile.write("\t{}\n".format(dsk))
    outFile.close()

def sigMerger(batchnum = 0,patterns = ["\d_0"]):
    mergedFile = PyPDF2.PdfMerger()
    tempIndexes = []
    batchcontents = os.listdir("outputData/Batch{}".format(batchnum))
    currentPage = 0
    for folder in batchcontents:
        subpath = "outputData/Batch{}/{}".format(batchnum,folder)
        for file in os.listdir(subpath):
            for pattern in patterns:
                if re.search(pattern,file) and ".pdf" in file:
                    filepath = "{}/{}".format(subpath,file)
                    filepages = (currentPage,currentPage + len(PyPDF2.PdfReader(filepath).pages))
                    tempIndexes.append((filepath,filepages))
                    mergedFile.append(filepath)
                    currentPage += len(PyPDF2.PdfReader(filepath).pages)
    mergedFile.write("outputData/Sigmerger.pdf")
    os.startfile("{}/outputData/Sigmerger.pdf".format(abspath))
    tempIndexFile = open("outputData/Sigindex.txt","w")
    for line in tempIndexes:
        tempIndexFile.write(line[0] + "|" + str(line[1][0]) + "|" + str(line[1][1]) + "\n")
    tempIndexFile.close()
    input("Press enter when ready to split:")
    sigSplitter("outputData/Sigmerger.pdf",tempIndexes)

def sigSplitter(sigfile,indexList):
    sigReader = PyPDF2.PdfReader(sigfile)
    for entry in indexList:
        tempWriter = PyPDF2.PdfWriter()
        tempWriter.append(sigReader,entry[1])
        tempWriter.write(entry[0])
            


def dictCheck(inData):
    #Check dictionary to make absolutely sure dates are mmddyy with no extraneous things
    for key in inData.keys():
        if "Dt" in key:
            if not inData[key].isnumeric():
                print("Thing found: {} {}".format(key,inData[key]))
            elif len(inData[key]) != 6:
                print("Thing found: {} {}".format(key,inData[key]))

def findexts():
    #Finds filed extensions in the filedpath
    #This gets a list of everything with the extension focuments in it
    numlist = []
    matchindex = 0
    for root,dir_names,file_names in os.walk(filedpath):
        if len(dir_names) == 0:
            #Looking for: motion, order, decs
            matches = 0
            matchlist = ["^00 .+ Motion.pdf$","^01 .+ Order.pdf$","^02 .+ Dec re Ext.pdf$","^03 .+ Dec re Int.pdf$"]
            for file in file_names:
                for pattern in matchlist:
                    if re.search(pattern,file):
                        matches += 1
            order = ""
            for file in file_names:
                if re.search(matchlist[0],file):
                    order = file
            if matches == 4:
                print("Match {}".format(matchindex))
                matchindex += 1
                tempData = OCRrun.getInfo("{}/{}".format(root,file),0,10)
                numlist.append(tempData["Legal"])
    return numlist

def findKeyword(pattern):
    #Searches folder names in filedpath for a pattern to help find old filed cases when they're rejected and need resubmission
    for root,dirnames,filenames in os.walk(filedpath):
        if pattern in root:
            print(root)

def processMailnos(inFile,skips = []):
    #Gets mail tracking numbers from a manually created txt file and formats them for automatic entry into PAMAR
    outFile = open("inputData/mailnosRun.txt","w")
    inData = open("inputData/{}".format(inFile),"r").read().splitlines()
    for x in range(int(len(inData) / 3)):
        if inData[3*x] not in skips: outFile.write("{}|poe {}|db {}\n".format(inData[3*x],inData[3*x+1],inData[3*x+2]))
    outFile.close()

def newcaseChopper(inFile,batchnum):
    #Simple chopper for new cases, which consist of a 2-page summons, a 1-page coversheet, and a variable-length (4+ page) complaint
    #Cases are given an index number and legal number when written to file as separate documents
    indexData = OCRrun.newCaseSplit(inFile)
    newCaseReader = PyPDF2.PdfReader("inputData/{}".format(inFile))
    choplist = [index[2] for index in indexData]
    numlist = [index[1] for index in indexData]
    choplist.append(len(newCaseReader.pages))
    for x in range(int(len(numlist) / 2)):
        caseno = numlist[2*x]
        summonsrange = (choplist[2*x],choplist[2*x + 1] - 1)
        coverrange = (choplist[2*x + 1] - 1,choplist[2*x + 1])
        complaintrange = (choplist[2*x+1],choplist[2*x+2])
        summons = PyPDF2.PdfWriter()
        cover = PyPDF2.PdfWriter()
        complaint = PyPDF2.PdfWriter()
        summons.append(newCaseReader,summonsrange)
        cover.append(newCaseReader,coverrange)
        complaint.append(newCaseReader,complaintrange)
        summons.write("outputData/Batch{}/{} {} Summons.pdf".format(batchnum,x,caseno))
        cover.write("outputData/Batch{}/{} {} Cover.pdf".format(batchnum,x,caseno))
        complaint.write("outputData/Batch{}/{} {} Complaint.pdf".format(batchnum,x,caseno))

def fastparse(infile,pattern):
    #Quick parser to avoid having to debug DataProcessor for a simple thing, now obsolete
    inData = open(infile,"r").read().splitlines()
    outData = []
    for line in inData:
        if re.search(pattern,line):
            outData.append(line)
    return outData

def greenmailbacks():
    #Create scanner backs for mail returns
    sublist = []
    for line in fastparse("inputData/greenmail.txt","L\d+ P\d+"):
        sublist.append(line.split("P")[-1])
    outdoc = docx.Document()
    style = outdoc.styles["Normal"]
    style.font.name = "Courier"
    style.font.size = docx.shared.Pt(20)
    for case in sublist:
        numpar = outdoc.add_paragraph()
        numpar.style = outdoc.styles["Normal"]
        numpar.add_run("P{}".format(case))
        numpar.add_run().add_break(break_type = docx.enum.text.WD_BREAK.PAGE)
    outdoc.save("Greenmailbacks.docx")

def WDSchopper(filename,data):
    #Old simple chopper for withdrawals and substitutions
    guessData = OCRrun.docIndexes(filename,mode = 10,width = 2)
    fileData = PyPDF2.PdfReader("inputData/{}".format(filename))
    for x in range(len(guessData)):
        print("Loop {}".format(x))
        outFile = PyPDF2.PdfMerger()
        filename = "{}".format(x)
        for case in data:
            if case["Legal"] == guessData[x]:
                filename = case["FolderName"]
        outFile.append(fileData,pages = (2*x,2*x+2))
        outFile.write("outputData/00 {} WDS.pdf".format(filename))

def configs(mode):
    #Helper function to get possible valid sets of files
    if mode == 0:
        #Orders mode
        #Any Order or Special Order will have a file for all indexes indicated here
        configs = [["00","01","02","03","04_0","08","09_0","09_1"]]
        #An Order or Special Order will either have all files in one or all files in the other for each of these three possible cases
        configsets = [[["04_1"],["05_0","05_1"]],[["04_2"],["06_0","06_1"]],[["07_0"],[]]]
        for layer in configsets:
            newconfigs = []
            for x in layer:
                for y in configs:
                    newconfigs.append(y + x)
            configs = newconfigs
    elif mode == 1:
        #Extensions are much simpler, all should have exactly these four/five documents
        configs = [["00","01","02","03"],["00","01","02","03","04"]]
    return configs

def lens(mode):
    #Helper function to get expected lengths of the indicated file index for validation purposes
    if mode == 0:
        lendict = {"00" : [2],"01" : [2],"02":[1,2,3,4,5,6,7,8],"03":[1,2,3,4,5,6,7,8],"04_0":[2,3,4],"04_1":[1],"04_2":[1],"05_0":[2,4],"05_1":[2],"06_0":[2,4],
                   "06_1":[2],"07_0":[2],"08":[1],"09_0":[4,6,8,10,12,14,16,18],"09_1":[2,4,6,8,10,12,14,16]}
    elif mode == 1:
        lendict = {"00":[2],"01":[3],"02":[2],"03":[2],"04":[2,3]}
    return lendict

def validate(mode,batchnum):
    #Mode 0  validates regular orders, mode 1 validates extensions
    path = "outputData/Batch{}".format(batchnum)
    configList = configs(mode)
    lenList = lens(mode)
    for case in os.listdir(path):
        subpath = "{}/{}".format(path,case)
        sublist = os.listdir(subpath)
        indexlist = []
        for file in sublist:
            if ".pdf" in file:
                fileLen = len(PyPDF2.PdfReader("{}/{}".format(subpath,file)).pages)
                if fileLen not in lenList[file.split(" ")[0].replace(".pdf","")]:
                    print("Error: file {}/{} wrong length".format(subpath,file))
                indexlist.append(file.split(" ")[0].replace(".pdf",""))
        configFound = False
        for x in range(len(configList)):
            validConfig = True
            for prefix in configList[x]:
                if prefix not in indexlist:
                    validConfig = False
            if validConfig:
                configFound = True
        if not configFound:
            print("Error: file {} missing documents".format(subpath))

def addrcheck(batchnum,data):
    #Opens certs to double-check addresses, since the PAMAR output isn't always consistent in where things were sent
    #Address keys: AddrD, OrigAddr, AddrG
    for x,case in enumerate(data):
        casepath = "outputData/Batch{}/{}".format(batchnum,case["FolderName"])
        casefiles = os.listdir(casepath)
        if case["AddrD"] != case["OrigAddr"]:
            for file in casefiles:
                if file[:2] in ["04","05","06"]:
                    os.startfile("{}/{}/{}".format(abspath,casepath,file))
            for key in ["MailnoD","AddrD","OrigAddr","MailnoG","AddrG"]:
                print(key,case[key])
            input("{}: Press enter to continue:".format(x))
        else:
            print("{}: No error at {}".format(x,case["FolderName"]))

def choplist(inFile,batchnum):
    return [a for a in DataProcessor2.newParser(inFile) if a["FolderName"] in os.listdir("outputData/Batch{}".format(batchnum))]

def ABClist(variants):
    if 0 in variants: OCRrun.ABCsheet("Legal Documents","Snohomish","Yes","","0")
    if 1 in variants: OCRrun.ABCsheet("Legal Documents","Snohomish","","South","1")
    if 2 in variants: OCRrun.ABCsheet("Legal Documents","Snohomish","","Evergreen","2")
    if 3 in variants: OCRrun.ABCsheet("Legal Documents","Snohomish","","Everett","3")
    if 4 in variants: OCRrun.ABCsheet("Legal Documents","Snohomish","","Cascade","4")
    if 5 in variants: OCRrun.ABCsheet("Legal Documents","Pierce","","#1","5")
    if 6 in variants: OCRrun.ABCsheet("Legal Documents","Pierce","Yes","","6")
    outDoc = PyPDF2.PdfMerger()
    for x in range(7):
        if "outSheet{}.pdf".format(x) in os.listdir():
            outDoc.append("outSheet{}.pdf".format(x))
    outDoc.write("ABC.pdf")
    outDoc.close()
    for x in range(7):
        if "outSheet{}.pdf".format(x) in os.listdir():
            os.remove("outSheet{}.pdf".format(x))

def reext(data,batchnum = 0,batched = True):
    path = "outputData/Batch{}".format(batchnum)
    massWDS = PyPDF2.PdfMerger()
    for file in data:
        file["DtDc"] = "083023"
        file["Me"] = "Nathaniel Slivka"
        if batched:
            if file["FolderName"] in os.listdir(path):
                indexes = [a[:2] for a in os.listdir("{}/{}".format(path,file["FolderName"]))]
                DeclarationMaker2.makeDeclaration(file,"19","{}/{}/04 {} WDS.docx".format(path,file["FolderName"],file["case#"]))
                docx2pdf.convert("{}/{}/04 {} WDS.docx".format(path,file["FolderName"],file["case#"]),"{}/{}/04 {} WDS.pdf".format(path,file["FolderName"],file["case#"]))
                if file["County"] == "KING":
                    massWDS.append("{}/{}/04 {} WDS.pdf".format(path,file["FolderName"],file["case#"]),pages = (0,2))
        else:
            DeclarationMaker2.makeDeclaration(file,"19","outputData/{} WDS.docx".format(file["case#"]))
            docx2pdf.convert("outputData/{} WDS.docx".format(file["case#"]),"outputData/{} WDS.pdf".format(file["case#"]))
            massWDS.append("outputData/{} WDS.pdf".format(file["case#"]),pages = (0,2))
    massWDS.write("YetMoreWDS.pdf")

def SSNprocess(inData,inFile = "SSNs.txt"):
    fileData = open("outputData/{}".format(inFile),"r").read().splitlines()
    newOut = open("outputData/NewSSNs.txt","w")
    plist = [[a["FolderName"],a["Packet"]] for a in inData if a["County"] != "KING"]
    pindex = 0
    for x in range(len(fileData)):
        line = fileData[x]
        if "In case " in line:
            pnum = line.split("P")[-1]
            plist[pindex].append(pnum)
            newOut.write("In case {} {}\n".format(plist[pindex][0],plist[pindex][1]))
            pindex += 1
        else:
            newOut.write("{}\n".format(line))
    newOut.close()
    for line in plist:
        print(line)

def fastMOsplitter(inData,inDoc,batchnum,debug = False):
    tempReader = PyPDF2.PdfReader("inputData/{}".format(inDoc))
    doclen = len(tempReader.pages)
    starttime = time()
    if debug: print("Entering contentslist at {}".format(starttime))
    contentslist = []
    for a in range(doclen):
        inttime = time()
        contentslist.append([OCRrun.getInfo("inputData/{}".format(inDoc),a,12)["Legal"],a])
        print("Page {} done in {}".format(a,time() - inttime))
    endtime = time()
    if debug: print("Ending contentslist at {} after {}".format(endtime,endtime - starttime))
    uniquelist = {}
    for a in contentslist: uniquelist = Utilities.dictAppend(uniquelist,a[0],a[1])
    for a in uniquelist.keys():
        for b in [["00 {} Motion.pdf",0,1],["01 {} Order.pdf",2,4]]:
            filename = "outputData/Batch{}/{}/{}".format(batchnum,inData[a]["FolderName"],b[0].format(inData[a]["case#"]))
            print("Writing pages {} to {} to file {}".format(uniquelist[a][b[1]],uniquelist[a][b[2]],filename))
            if not debug:
                tempmerger = PyPDF2.PdfMerger()
                tempmerger.append(tempReader,pages = (int(uniquelist[a][b[1]]),int(uniquelist[a][b[2]]) + 1))
                tempmerger.write(filename)

def quickMerge(filePath,filePatterns):
    mergeFolderName = filePath.split("/")[-1]
    mergeFolderContents = os.listdir(filePath)
    fileMerger = PyPDF2.PdfMerger()
    for file in mergeFolderContents:
        if ".pdf" in file and file.split(" ")[0] in filePatterns:
            fileMerger.append("{}/{}".format(filePath,file))
    fileMerger.write("outputData/{} Merged.pdf".format(mergeFolderName))
    
def copytoPacket(batchnum,pattern):
    quickpath = "outputData/Batch{}".format(batchnum)
    copypath = "P:/debtors/ocr-pkt-input"
    for subfolder in os.listdir(quickpath):
        for subfile in os.listdir("{}/{}".format(quickpath,subfolder)):
            if pattern in subfile:
                shutil.copy("{}/{}/{}".format(quickpath,subfolder,subfile),"{}/{}".format(copypath,subfile))

def assembleResublist(patternlist):
    #There are *still* things that idea won't catch. Bastards.
    guesslist = {a:[] for a in range(len(patternlist))}
    for root,dirnames,filenames in os.walk(filedpath):
        subpath = root.split("Filed Documents")[1]
        subpathlist = subpath.split("/")
        #print(subpath)
        #What are we expecting to find here? Date/batch/name, or date/batch/batch/name, or date/name.
        #First: Check candidates to see if they contain a pattern.
        for a in range(len(patternlist)):
            #We're iterating through one path at a time.
            #Let's assume for simplicity that patternlist[0] is case numbers and patternlist[1] is names.
            for b in subpathlist:
                #We don't know which folder is the identifier
                if patternlist[a][0] in b:
                    #If the case number's there, great, we're done! yay!
                    guesslist[a].append(root)
                elif patternlist[a][1] in b and b.isalpha():
                    #If the name is there but not the case number it might be okay if it's just a name, but not if it's a different case number
                    guesslist[a].append(root)
                #At the moment I'm *pretty* sure that should catch everything.
    return guesslist
                


data0 = DataProcessor2.parseWrapper("092023 Ords Capture.txt",0)
#makeQuickFolders(data0)
lnumlist = ["{}|{}".format(a["Legal"],a["garns"]) for a in data0]

data1 = DataProcessor2.parseWrapper("092023 Exts Capture.txt",1)

data5 = DataProcessor2.parseWrapper("092723 Ext Precapture.txt",5)
data6 = DataProcessor2.newParser("092723 Ords Capture.txt")

#KeyInput2.advancedRecord("092023 New Garns Lnums.txt",instructionlist = ["$1","","2","11","n","orange","n","","","","","","","d","2","d","1","e","5","d","2","e","e","e"],isCapture = True)
#KeyInput2.advancedRecord("092023 New Garns Lnums.txt",instructionlist = ["$0","i","e","e"],isCapture = True)

test1 = DataProcessor2.packetParser("092023 New Garn Test Cap 1.txt")
test2 = DataProcessor2.newParser("092023 New Garn Test Cap 2.txt")
testdict1 = {a["Packet"]:a for a in test1}
testdict2 = {a["Packet"]:a for a in test2}

#Got some data, now what do we actually need to enter?
#legal, g, n, e, "", name1, name2, employer1, employer2, employeraddr1, employercsz, name of actual garnishee, ssn of garnishee, file fee ($12.00),
#service fee ($9.01), "", atty fee (10% of balance owing, from 100 to 300), c/mail (10.90), parte (by county), e, y, "", "", e, e

#How do we get these?
#print(test1[0].keys())
#print(test2[0].keys())
#Legal: test2["Legal"]
#Names? [b.split(", ")[1] + " " + b.split(", ")[0] for b in [a.strip() for a in test2["topdefs"] if a.strip() not in ["JOHN DOE","JANE DOE","DOE, JOHN","DOE, JANE"]]]
#Employer? Depends on who we're garnishing...

#OCRrun.gensplitter(data1,"092023 Exts MOs.pdf")

#PdfHandler2.countyGroup(data0)
#PdfHandler2.countyGroup(data1,batchnum = 1)
#OCRrun.newGensplitter(data1,"092023 Exts WDS.pdf",1,batchnum = 1)
#PdfHandler2.countyGroup(data1,batchnum = 1)

#errorCleanup(mode = 5,batchnum = 1)